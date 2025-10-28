"""
DOCX (Word document) reader implementation.

Auto-registers for Word MIME types.
"""

from docx import Document as DocxDocument
from .base_reader import BaseReader
from ..document import Document, TableData
from ..file_loader import reader_registry


@reader_registry.register(
    'application/vnd.openxmlformats-officedocument.wordprocessingml.document',  # docx
    'application/msword'  # doc
)
class DocxReader(BaseReader):
    """
    Word document reader.
    
    Extracts text, tables, and metadata from DOCX files.
    Auto-registered for Word MIME types.
    """
    
    def read(self, file_path: str, document: Document) -> None:
        """
        Read DOCX file and populate document.
        
        Args:
            file_path: Path to DOCX file
            document: Document object to populate
        """
        try:
            docx = DocxDocument(file_path)
            
            # Extract metadata
            core_props = docx.core_properties
            document.metadata['title'] = core_props.title or ''
            document.metadata['author'] = core_props.author or ''
            document.metadata['subject'] = core_props.subject or ''
            document.metadata['keywords'] = core_props.keywords or ''
            document.metadata['created'] = core_props.created.isoformat() if core_props.created else ''
            document.metadata['modified'] = core_props.modified.isoformat() if core_props.modified else ''
            
            # Extract text from paragraphs
            text_content = []
            for para in docx.paragraphs:
                if para.text.strip():
                    text_content.append(para.text)
            
            document.raw_text = '\n\n'.join(text_content)
            
            # Extract tables
            for table_idx, table in enumerate(docx.tables):
                # Extract headers (first row)
                headers = [cell.text.strip() for cell in table.rows[0].cells]
                
                # Extract data rows
                rows = []
                for row in table.rows[1:]:
                    row_data = [cell.text.strip() for cell in row.cells]
                    rows.append(row_data)
                
                if rows:  # Only add if there's data
                    table_data = TableData(
                        headers=headers,
                        rows=rows,
                        metadata={
                            'table_index': table_idx,
                            'row_count': len(rows),
                            'column_count': len(headers)
                        }
                    )
                    document.add_table(table_data)
            
            # Store document statistics
            document.metadata['paragraph_count'] = len(docx.paragraphs)
            document.metadata['table_count'] = len(docx.tables)
            document.structure.page_count = 1  # DOCX doesn't have fixed pages
            
            if not document.raw_text.strip():
                document.add_warning("DOCX contains no text content")
            
        except Exception as e:
            document.add_error(f"Failed to read DOCX: {str(e)}")
            raise